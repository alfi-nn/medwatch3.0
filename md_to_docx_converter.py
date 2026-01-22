import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pathlib import Path

SOURCE_FILE = r"C:\Users\kjalf\.gemini\antigravity\brain\1c73620b-3c53-4296-b532-382081ee3755\hgt_deep_dive.md"
OUTPUT_FILE = "HGT_Architecture_Deep_Dive.docx"

def add_formatted_text(paragraph, text):
    """Parses simple markdown (bold, italic, code) and adds runs."""
    # Split by bold markers **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for code ` `
            code_parts = re.split(r'(`.*?`)', part)
            for cpart in code_parts:
                if cpart.startswith('`') and cpart.endswith('`'):
                    run = paragraph.add_run(cpart[1:-1])
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(200, 50, 50) 
                else:
                    paragraph.add_run(cpart)

def convert_md_to_docx():
    doc = docx.Document()
    
    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    try:
        with open(SOURCE_FILE, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: Could not find source file at {SOURCE_FILE}")
        return

    in_code_block = False
    
    for line in lines:
        line = line.strip()
        
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Lists
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        elif line.startswith('1. '):
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, line[3:])
            
        # Normal Text
        elif line:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    doc.save(OUTPUT_FILE)
    print(f"Successfully created {OUTPUT_FILE}")

if __name__ == "__main__":
    convert_md_to_docx()
